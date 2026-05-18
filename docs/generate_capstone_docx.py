"""Generate SiglaCast ITP capstone documentation (DOCX)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.expanduser(
    r"~\Downloads\SiglaCast_ITP_Documentation.docx"
)


def set_margins(doc):
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def style_paragraph(p, bold=False, italic=False, size=11, align=None, space_after=0):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        run.italic = italic


def add_text(doc, text, bold=False, italic=False, align=None, space_after=0):
    p = doc.add_paragraph(text)
    style_paragraph(p, bold=bold, italic=italic, align=align, space_after=space_after)
    return p


def add_chapter(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run(f"CHAPTER {number}\n")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r2 = p.add_run(title.upper())
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(11)


def add_section(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(f"{number} {title}")
    r.bold = True
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(11)


def add_body(doc, text):
    for para in text.strip().split("\n\n"):
        add_text(doc, para.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def main():
    doc = Document()
    set_margins(doc)

    # --- TITLE PAGE ---
    for line in [
        "DAVAO ORIENTAL STATE UNIVERSITY",
        "Guang-guang, Dahican, City of Mati, Davao Oriental",
        "",
        "SIGLACAST: VOTING AND COMMUNITY APPLICATION",
        "FOR DAVAO ORIENTAL STATE UNIVERSITY EVENTS",
        "",
        "Presented to the",
        "Bachelor of Science in Information Technology",
        "Davao Oriental State University",
        "Panabo City, Davao del Norte",
        "",
        "In Partial Fulfillment",
        "of the Requirements for the Course",
        "ITP 121 – INTEGRATIVE PROGRAMMING AND TECHNOLOGIES 1",
        "",
        "Proponent 1",
        "Proponent 2",
        "Proponent 3",
        "Proponent 4",
        "Proponent 5",
        "",
        "MAY 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if "SIGLACAST" in line:
                run.bold = True

    doc.add_page_break()

    # --- TABLE OF CONTENTS (manual per template) ---
    add_text(doc, "TABLE OF CONTENTS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    toc_lines = [
        "TITLE PAGE ................................................................. i",
        "TABLE OF CONTENTS .......................................................... ii",
        "CHAPTER",
        "1 INTRODUCTION",
        "    1.1 Rationale .............................................................. 1",
        "    1.2 Purpose and Description .............................................. 1",
        "    1.3 General Objectives ..................................................... 2",
        "    1.4 Specific Objectives .................................................... 2",
        "    1.5 Scope and Limitations .................................................. 3",
        "    1.6 Significance of the Study .............................................. 3",
        "2 METHODOLOGY",
        "    2.1 Requirement Specification ............................................ 4",
        "        2.1.1 Functional Requirements ........................................ 4",
        "        2.1.2 Non-Functional Requirements .................................... 5",
        "    2.2 System Analysis ........................................................ 5",
        "        2.2.1 Business Process ............................................... 5",
        "        2.2.2 Use Case ......................................................... 6",
        "        2.2.3 System Architecture .............................................. 7",
        "        2.2.4 Integrative Programming Topics Implementation .................... 7",
        "        2.2.5 Data Storage and Message Broker Flow ............................. 8",
        "        2.2.6 Development Procedure ............................................ 9",
        "REFERENCES ................................................................. 10",
        "APPENDIX A – SYSTEM SCREENSHOTS ............................................ 11",
        "APPENDIX B – PHOTO DOCUMENTATION ............................................. 12",
    ]
    for line in toc_lines:
        add_text(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # --- CHAPTER 1 ---
    add_chapter(doc, 1, "INTRODUCTION")

    add_section(doc, "1.1", "Rationale")
    add_body(
        doc,
        """
University events at Davao Oriental State University, such as student council elections, campus polls, intramural awards, and organizational activities, require timely participation, transparent vote counting, and clear communication among students, faculty, and administrators. Traditional methods that rely on paper ballots, manual tally sheets, and scattered social media updates often result in delayed results, limited audit trails, and weak engagement outside the physical venue.

SiglaCast addresses these concerns by providing a unified web-based voting and community platform. The system allows authenticated students to participate in campus events, cast votes under configurable rules, view live tallies, read announcements, interact in a community feed, exchange private messages, and manage personal profiles. Administrators can create events, publish announcements, and monitor participation through role-based dashboards. The project also integrates core topics from Integrative Programming and Technologies 1, including object-oriented programming, messaging brokers, XML processing, XSLT transformation, and administrative scripting, so that the solution reflects both practical campus needs and course learning outcomes.
        """,
    )

    add_section(doc, "1.2", "Purpose and Description")
    add_body(
        doc,
        """
The purpose of SiglaCast is to design and implement an advanced information system that supports secure event voting and digital community engagement for Davao Oriental State University. The system is implemented as a full-stack web application composed of a React frontend, a Node.js and Express backend REST API, JSON-based persistent storage, optional RabbitMQ or Kafka message brokers, and admin automation scripts.

SiglaCast uses a blue, yellow, and white visual theme for clarity and brand consistency. Students register or sign in with institutional email credentials, browse open events, review candidates and rules, submit votes within per-event limits, and observe results through visual tally displays. The community module supports posts with images, reactions, and comments. The messaging module allows user search, friend connections, and private chat threads. Administrators manage events, announcements, and operational dashboards. XML export, XML parsing, and XSLT-based HTML reporting demonstrate structured data handling required by the integrative programming curriculum.
        """,
    )

    add_section(doc, "1.3", "General Objectives")
    add_body(
        doc,
        """
The general objective of this study is to develop SiglaCast, an integrative web-based voting and community application for Davao Oriental State University events that improves participation, transparency, and communication while demonstrating required programming technologies in a single cohesive system.
        """,
    )

    add_section(doc, "1.4", "Specific Objectives")
    objectives = [
        "To implement secure user authentication and role-based access for students and administrators using JSON Web Tokens and password hashing.",
        "To design and develop event management with configurable voting strategies, vote limits, candidate profiles, and live tally visualization.",
        "To provide community features including posts, image uploads, reactions, comments, announcements, and in-app notifications.",
        "To implement private messaging with user search, friend management, conversation listing, and real-time refresh through polling.",
        "To apply object-oriented programming principles including inheritance, encapsulation, and polymorphism in backend domain modeling.",
        "To integrate a messaging broker (RabbitMQ or Apache Kafka) for asynchronous event publishing on votes, posts, announcements, and messages.",
        "To implement XML generation, XML parsing, and XSLT transformation for event data reporting.",
        "To create administrative scripts for environment setup, queue monitoring, and system maintenance using scripting languages and advanced scripting techniques.",
    ]
    for i, obj in enumerate(objectives, 1):
        add_text(doc, f"{i}. {obj}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_section(doc, "1.5", "Scope and Limitations")
    add_body(
        doc,
        """
Scope. SiglaCast covers user registration and login, student and admin dashboards, event creation and voting, live tallies, community posts, announcements, notifications, profile and avatar management, private messaging, XML and XSLT endpoints, and message broker integration. The system targets university event scenarios such as elections and campus polls within the Davao Oriental State University context.

Limitations. The system uses a JSON file (db.json) as its primary data store rather than a relational database server, which may limit scalability under very high concurrent load. Private chat delivery uses HTTP polling instead of WebSockets or push notifications. Message broker consumers currently log events for demonstration rather than driving separate microservices. The application is designed for academic and pilot deployment; production hardening such as clustered hosting, formal penetration testing, and institutional single sign-on integration are outside the current scope.
        """,
    )

    add_section(doc, "1.6", "Significance of the Study")
    add_body(
        doc,
        """
SiglaCast benefits students by offering a convenient and transparent channel to vote and stay informed about campus events. Administrators gain centralized tools to publish events, enforce voting rules, and communicate announcements. Faculty and institutional stakeholders may use exported XML and HTML reports for documentation and review.

For the field of information technology education, the project demonstrates how integrative programming topics—OOP, messaging brokers, XML, XSLT, and scripting—can be applied in a real-world campus system rather than isolated laboratory exercises. The study may serve as a reference for future capstone projects that combine modern web development with enterprise integration patterns at Davao Oriental State University.
        """,
    )

    doc.add_page_break()

    # --- CHAPTER 2 ---
    add_chapter(doc, 2, "METHODOLOGY")

    add_section(doc, "2.1", "Requirement Specification")

    add_section(doc, "2.1.1", "Functional Requirements")
    functional = [
        "The system shall allow users to register, log in, log out, and refresh sessions using access and refresh tokens.",
        "The system shall distinguish student and admin roles and restrict admin-only routes such as event and announcement creation.",
        "The system shall allow administrators to create events with title, description, rules, cover image, candidates, voting strategy (single or weighted), status, and maximum votes per user.",
        "The system shall allow authenticated students to cast votes on open events and enforce per-event vote limits.",
        "The system shall compute and display live vote tallies with visual progress indicators.",
        "The system shall provide a community module for creating posts with optional images, reactions, and comments.",
        "The system shall support announcements and user-specific notifications.",
        "The system shall allow users to update profile name, password, and avatar image.",
        "The system shall support user search, adding friends, listing conversations, sending private messages, and marking messages as read.",
        "The system shall export events as XML, parse submitted XML, and transform XML to HTML using XSLT.",
        "The system shall publish asynchronous events to a message broker when votes are cast, posts are created, announcements are made, and messages are sent.",
    ]
    for req in functional:
        add_text(doc, f"• {req}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_section(doc, "2.1.2", "Non-Functional Requirements")
    non_functional = [
        "Usability: The interface shall use a consistent blue, yellow, and white theme with clear navigation across dashboard, events, community, messages, announcements, notifications, and profile modules.",
        "Security: Passwords shall be stored using bcrypt hashing; API routes shall require valid JWT access tokens except for registration, login, and refresh endpoints.",
        "Performance: The client shall poll for conversation and tally updates at reasonable intervals during active use.",
        "Maintainability: Backend code shall use modular classes for users, vote strategies, and message broker implementations.",
        "Portability: The system shall run on Node.js with optional Docker services for RabbitMQ and Kafka.",
        "Reliability: Data changes shall be persisted to db.json after successful operations; broker connection failures shall fall back to an in-memory broker.",
    ]
    for req in non_functional:
        add_text(doc, f"• {req}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_section(doc, "2.2", "System Analysis")

    add_section(doc, "2.2.1", "Business Process")
    add_body(
        doc,
        """
The SiglaCast business process begins when an administrator creates and opens a campus event. Students authenticate, review event details and candidates, and submit votes within allowed limits. The system records each vote, updates tallies, notifies the voter, and publishes a vote.cast broker event. Parallel processes allow students to post in the community feed, read announcements, receive notifications, and communicate through private messages. Administrators monitor dashboards, publish announcements, and may export event data through XML and XSLT endpoints for reporting. The process ends when an event is closed and final tallies are available for review.
        """,
    )

    add_section(doc, "2.2.2", "Use Case")
    use_cases = [
        "UC-01 Register Account – A new student creates an account with name, email, course, and password.",
        "UC-02 Login – A user authenticates and receives access and refresh tokens.",
        "UC-03 Vote in Event – A student selects a candidate and submits a vote on an open event.",
        "UC-04 View Live Tally – A user opens event details and views vote counts and percentage bars.",
        "UC-05 Create Event (Admin) – An administrator defines event metadata, candidates, and voting rules.",
        "UC-06 Create Community Post – A user publishes text and optional image content to the community feed.",
        "UC-07 React and Comment – A user reacts to or comments on an existing post.",
        "UC-08 Send Private Message – A user searches for another user, opens a chat thread, and sends a message.",
        "UC-09 Add Friend – A user sends a friend request relationship stored in the system.",
        "UC-10 Export XML Report – An authorized user requests XML or HTML transformation of event data.",
        "UC-11 Publish Announcement (Admin) – An administrator broadcasts campus-wide announcement text.",
    ]
    for uc in use_cases:
        add_text(doc, uc, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_section(doc, "2.2.3", "System Architecture")
    add_body(
        doc,
        """
SiglaCast follows a three-tier architecture. The presentation tier is a React single-page application built with Vite and React Router, communicating over HTTP with JSON payloads. The application tier is an Express.js REST API on port 4000 that handles authentication, business logic, file uploads, XML processing, and broker publishing. The data tier consists of a JSON file (backend/src/data/db.json) for structured records and a local uploads directory for images. An optional integration tier connects to RabbitMQ (default) or Apache Kafka through Docker Compose for asynchronous event queues named vote.cast, post.created, announcement.created, and message.sent.
        """,
    )

    add_section(doc, "2.2.4", "Integrative Programming Topics Implementation")
    add_body(
        doc,
        """
Object-Oriented Programming. The backend defines a User base class with private fields and getters, Student and Admin subclasses for inheritance, and VoteStrategy with SingleVoteStrategy and WeightedVoteStrategy subclasses for polymorphic tally computation. Encapsulation is applied through private class fields and controlled access patterns.

Messaging Broker. A MessageBroker abstraction supports RabbitBroker, KafkaBroker, and InMemoryBroker implementations. On startup, the server connects based on the BROKER environment variable and registers consumers that process published events.

XML and XML Parsing. The API exposes GET /api/xml/events.xml to generate event XML using xmlbuilder2 and POST /api/xml/parse to parse XML input using fast-xml-parser.

XSL and XSLT. The API exposes GET /api/xml/events.html, which loads events.xsl and transforms event XML into HTML using xslt-processor for browser-readable reports.

Scripting Languages. JavaScript is used for the backend server, frontend client, and Node.js administrative utilities.

Scripting for System Administration. PowerShell (setup-env.ps1) and Bash (setup-env.sh) scripts automate environment file creation and local setup.

Advanced Scripting Techniques. The monitor-queue.js script accepts CLI arguments, maintains a rolling history buffer, and refreshes terminal output for queue monitoring demonstration.
        """,
    )

    add_section(doc, "2.2.5", "Data Storage and Message Broker Flow")
    add_body(
        doc,
        """
Primary data is stored in db.json under collections: users, events, votes, posts, announcements, notifications, friends, and messages. When a user performs an action, the API updates the in-memory database object and calls saveDb() to write JSON to disk. Uploaded images are stored under backend/uploads with URLs referenced in JSON records.

For message broker flow, when a vote is recorded the API saves the vote, then publishes { userId, eventId, candidateId } to the vote.cast topic. When a post is created, the full post object is published to post.created. When a private message is sent, the message is saved to db.messages, a notification is created for the recipient, and a lightweight payload is published to message.sent. Broker consumers on the same server log these events; in a production extension, separate services could subscribe for email, analytics, or audit logging.
        """,
    )

    add_section(doc, "2.2.6", "Development Procedure")
    add_body(
        doc,
        """
The development procedure followed an iterative approach. Requirements were derived from the capstone concept and integrative programming topic list. The backend API was implemented first with authentication, events, voting, and data persistence. Frontend pages were developed with React and connected through a shared API service module. Community, messaging, profile, and notification features were added in subsequent iterations. Integrative topics were embedded in server.js and supporting scripts. Docker Compose was prepared for RabbitMQ and Kafka testing. The system was validated through manual testing with demo student and admin accounts (for example, ana@dorsu.edu.ph and admin@dorsu.edu.ph). Documentation and capstone manuscript preparation concluded the cycle.
        """,
    )

    doc.add_page_break()

    # --- REFERENCES ---
    add_text(doc, "REFERENCES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    references = [
        "Apache Kafka. (2024). Apache Kafka documentation. https://kafka.apache.org/documentation/",
        "Express.js. (2024). Express web framework for Node.js. https://expressjs.com/",
        "Fielding, R. T. (2000). Architectural styles and the design of network-based software architectures (Doctoral dissertation, University of California, Irvine).",
        "JSON Web Token (JWT). (2024). RFC 7519 – JSON Web Token (JWT). IETF.",
        "Node.js. (2024). Node.js JavaScript runtime. https://nodejs.org/",
        "RabbitMQ. (2024). RabbitMQ documentation. https://www.rabbitmq.com/documentation.html",
        "React. (2024). React – The library for web and native user interfaces. https://react.dev/",
        "W3C. (1999). XSL transformations (XSLT) version 1.0. https://www.w3.org/TR/xslt-10/",
        "World Wide Web Consortium. (2008). Extensible Markup Language (XML) 1.0 (Fifth Edition). https://www.w3.org/TR/xml/",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    doc.add_page_break()

    add_text(doc, "APPENDIX A", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "SYSTEM SCREENSHOTS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_body(
        doc,
        """
Insert screenshots of the following SiglaCast modules in this appendix: (1) Login and registration page, (2) Student dashboard, (3) Admin dashboard, (4) Events list and event detail with live tally, (5) Community feed with post and reactions, (6) Messages page with search and chat thread, (7) Announcements and notifications pages, (8) Profile page with avatar upload, and (9) XML/HTML export sample in a browser. Label each figure as Figure A-1, Figure A-2, and so forth according to department guidelines.
        """,
    )

    doc.add_page_break()

    add_text(doc, "APPENDIX B", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "PHOTO DOCUMENTATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_body(
        doc,
        """
Insert photo documentation of the development process in this appendix, including team meetings, system testing, defense preparation, and presentation setup. Include captions, dates, and participant names as required by the instructor.
        """,
    )

    os.makedirs(os.path.dirname(OUTPUT) or ".", exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
